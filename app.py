"""
SentimentML — Multi-Format Sentiment Analyzer (Streamlit Edition)
Supports: CSV, PDF, DOCX, TXT, XLSX, JSON
Uses TF-IDF + Logistic Regression for structured data,
and keyword/heuristic NLP for unstructured text files.

Run with:  streamlit run app.py
"""

import io
import re
import string
import json

import numpy as np
import pandas as pd
import streamlit as st
import plotly.graph_objects as go

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, classification_report
from sklearn.pipeline import Pipeline

import nltk
from nltk.corpus import stopwords
from nltk.stem import SnowballStemmer

import PyPDF2
import docx as python_docx


# ────────────────────────────────────────────────────────────────
# Page setup
# ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="SentimentML — Multi-Format Analyzer",
    page_icon="🧠",
    layout="wide",
)

ALLOWED_EXTENSIONS = {"csv", "pdf", "docx", "txt", "xlsx", "xls", "json"}


@st.cache_resource(show_spinner=False)
def _load_nltk():
    nltk.download("stopwords", quiet=True)
    nltk.download("punkt", quiet=True)
    return SnowballStemmer("english"), set(stopwords.words("english"))


_stemmer, _stopwords = _load_nltk()


# ────────────────────────────────────────────────────────────────
# Text cleaning / labeling
# ────────────────────────────────────────────────────────────────
def clean_text(text):
    text = str(text).lower()
    text = re.sub(r"\[.*?\]", "", text)
    text = re.sub(r"https?://\S+|www\.\S+", "", text)
    text = re.sub(r"<.*?>", "", text)
    text = re.sub(r"[%s]" % re.escape(string.punctuation), "", text)
    text = re.sub(r"\n", " ", text)
    text = re.sub(r"\w*\d\w*", "", text)
    tokens = [w for w in text.split() if w not in _stopwords]
    tokens = [_stemmer.stem(w) for w in tokens]
    return " ".join(tokens)


def normalize_label(val):
    v = str(val).strip().lower()
    if v in ("positive", "pos", "good", "happy", "joy", "love", "like",
              "very positive", "5", "4"):
        return "Positive"
    if v in ("negative", "neg", "bad", "sad", "anger", "hate", "dislike",
              "very negative", "1", "2"):
        return "Negative"
    return "Neutral"


# ────────────────────────────────────────────────────────────────
# Keyword-based heuristic sentiment (unstructured docs / live text)
# ────────────────────────────────────────────────────────────────
POS_WORDS = {
    "good", "great", "excellent", "amazing", "wonderful", "best", "love", "perfect",
    "awesome", "fantastic", "happy", "beautiful", "nice", "superb", "outstanding",
    "brilliant", "incredible", "fabulous", "splendid", "delightful", "positive",
    "pleased", "satisfied", "enjoy", "enjoyed", "recommend", "helpful", "impressed",
    "quality", "reliable", "fast", "easy", "clean", "friendly", "professional",
    "comfortable", "pleasant", "fun", "grateful", "glad", "thrilled", "excited",
    "hopeful", "joy", "joyful", "cheerful", "optimistic", "smooth", "efficient",
    "innovative", "skilled", "experienced", "dedicated", "passionate", "creative",
    "strong", "capable", "proficient", "expert", "accomplished", "achievement",
}
NEG_WORDS = {
    "bad", "terrible", "awful", "worst", "hate", "poor", "horrible", "disappointing",
    "useless", "failure", "broken", "waste", "ugly", "wrong", "never", "problem",
    "issue", "slow", "expensive", "rude", "unfriendly", "dirty", "disgusting",
    "annoying", "frustrating", "unreliable", "difficult", "confusing", "unhelpful",
    "negative", "sad", "angry", "upset", "dissatisfied", "regret", "avoid",
    "damaged", "defective", "faulty", "inferior", "mediocre", "pathetic", "absurd",
    "unacceptable", "misleading", "fraud", "scam", "fake", "lie", "error", "bug",
    "crash", "broken", "failed", "miss", "lack", "weak", "poor", "limited",
}
NEG_AMP = {"not", "no", "never", "neither", "nor", "hardly", "barely", "scarcely"}


def heuristic_sentiment(text):
    words = re.findall(r"\b\w+\b", str(text).lower())
    pos = neg = 0
    for i, w in enumerate(words):
        negate = i > 0 and words[i - 1] in NEG_AMP
        if w in POS_WORDS:
            neg += 1.2 if negate else 0
            pos += 0 if negate else 1
        elif w in NEG_WORDS:
            pos += 0.4 if negate else 0
            neg += 0 if negate else 1
    if pos > neg * 1.1:
        return "Positive"
    if neg > pos * 1.1:
        return "Negative"
    return "Neutral"


def predict_live_text(raw):
    """Sentence-level heuristic prediction used by the 'live text' box."""
    sentences = re.split(r"(?<=[.!?])\s+", raw)
    if len(sentences) == 1:
        sentences = [raw]

    sentence_results = []
    total_pos = total_neg = 0

    for sent in sentences:
        words = re.findall(r"\b\w+\b", sent.lower())
        pos = neg = 0
        for i, w in enumerate(words):
            negate = i > 0 and words[i - 1] in NEG_AMP
            if w in POS_WORDS:
                neg += 1.5 if negate else 0
                pos += 0 if negate else 1
            elif w in NEG_WORDS:
                pos += 0.5 if negate else 0
                neg += 0 if negate else 1
            if w in ("very", "really", "extremely", "absolutely", "totally", "so", "quite") and i + 1 < len(words):
                nw = words[i + 1]
                if nw in POS_WORDS:
                    pos += 0.5
                if nw in NEG_WORDS:
                    neg += 0.5

        pos += sent.count("!") * 0.3
        if sent.isupper() and len(sent) > 4:
            pos += 0.3
            neg += 0.3

        total_pos += pos
        total_neg += neg

        if pos > neg:
            label = "Positive"
        elif neg > pos:
            label = "Negative"
        else:
            label = "Neutral"

        sentence_results.append({
            "sentence": sent.strip(), "label": label,
            "pos": round(pos, 2), "neg": round(neg, 2),
        })

    total = total_pos + total_neg
    if total == 0:
        pos_pct = neu_pct = neg_pct = 33.3
        overall = "Neutral"
    else:
        sm = 0.05
        rp = total_pos / total
        rn = total_neg / total
        ru = max(0, 1 - rp - rn)
        d = rp + rn + ru + 3 * sm
        pos_pct = round((rp + sm) / d * 100, 1)
        neg_pct = round((rn + sm) / d * 100, 1)
        neu_pct = round(100 - pos_pct - neg_pct, 1)
        if total_pos > total_neg * 1.2:
            overall = "Positive"
        elif total_neg > total_pos * 1.2:
            overall = "Negative"
        else:
            overall = "Neutral"

    emoji = {"Positive": "😊", "Negative": "😞", "Neutral": "😐"}[overall]

    return {
        "overall": overall, "emoji": emoji,
        "pos_pct": pos_pct, "neg_pct": neg_pct, "neu_pct": neu_pct,
        "sentences": sentence_results,
        "word_count": len(re.findall(r"\b\w+\b", raw)),
    }


# ────────────────────────────────────────────────────────────────
# File parsers
# ────────────────────────────────────────────────────────────────
def parse_csv(file_bytes):
    try:
        df = pd.read_csv(io.StringIO(file_bytes.decode("utf-8", errors="ignore")))
        return df, None
    except Exception as e:
        return None, f"CSV parse error: {e}"


def parse_xlsx(file_bytes):
    try:
        df = pd.read_excel(io.BytesIO(file_bytes))
        return df, None
    except Exception as e:
        return None, f"Excel parse error: {e}"


def parse_json(file_bytes):
    try:
        data = json.loads(file_bytes.decode("utf-8", errors="ignore"))
        if isinstance(data, list):
            df = pd.DataFrame(data)
        elif isinstance(data, dict):
            df = pd.DataFrame([data]) if any(not isinstance(v, list) for v in data.values()) else pd.DataFrame(data)
        else:
            return None, "JSON must be an array of objects"
        return df, None
    except Exception as e:
        return None, f"JSON parse error: {e}"


def parse_pdf(file_bytes):
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        all_text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                all_text += t + "\n"
        chunks = re.split(r"(?<=[.!?])\s+|\n{2,}", all_text)
        chunks = [c.strip() for c in chunks if len(c.strip()) > 15]
        if not chunks:
            return None, "No readable text found in PDF"
        df = pd.DataFrame({"review": chunks})
        return df, None
    except Exception as e:
        return None, f"PDF parse error: {e}"


def parse_docx(file_bytes):
    try:
        doc = python_docx.Document(io.BytesIO(file_bytes))
        texts = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if len(t) > 10:
                texts.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if len(t) > 10:
                        texts.append(t)
        if not texts:
            return None, "No readable text found in DOCX"
        df = pd.DataFrame({"review": texts})
        return df, None
    except Exception as e:
        return None, f"DOCX parse error: {e}"


def parse_txt(file_bytes):
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        lines = [l.strip() for l in text.splitlines() if len(l.strip()) > 10]
        if not lines:
            return None, "No readable text found in TXT file"
        df = pd.DataFrame({"review": lines})
        return df, None
    except Exception as e:
        return None, f"TXT parse error: {e}"


def parse_file(filename, file_bytes):
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    structured = ext in ("csv", "xlsx", "xls", "json")
    if ext == "csv":
        result = parse_csv(file_bytes)
    elif ext in ("xlsx", "xls"):
        result = parse_xlsx(file_bytes)
    elif ext == "json":
        result = parse_json(file_bytes)
    elif ext == "pdf":
        result = parse_pdf(file_bytes)
    elif ext == "docx":
        result = parse_docx(file_bytes)
    elif ext == "txt":
        result = parse_txt(file_bytes)
    else:
        result = (None, f"Unsupported file type: .{ext}")
    return result, structured, ext


# ────────────────────────────────────────────────────────────────
# Core train + predict pipeline (mirrors original /predict route)
# ────────────────────────────────────────────────────────────────
def run_pipeline(filename, file_bytes):
    (df, parse_error), is_structured, ext = parse_file(filename, file_bytes)
    if parse_error:
        return {"error": parse_error}

    cols = [str(c).strip() for c in df.columns]
    df.columns = cols

    review_col = (
        next((c for c in cols if c.lower() in ["review", "text", "comment", "clean_comment",
                                                "tweet", "feedback", "description", "message",
                                                "content", "body"]), None)
        or next((c for c in cols if any(k in c.lower() for k in
                 ["review", "comment", "text", "tweet", "feedback", "content"])), None)
        or next((c for c in cols if df[c].dtype == object), None)
    )
    rating_col = (
        next((c for c in cols if c.lower() in ["rating", "sentiment", "label", "category",
                                                "class", "score", "polarity"]), None)
        or next((c for c in cols if any(k in c.lower() for k in
                 ["rating", "sentiment", "label", "categor", "class", "score"])), None)
    )
    product_col = next((c for c in cols if "product" in c.lower()), None)

    if not review_col:
        return {"error": f"No text column found. Detected columns: {', '.join(cols)}"}

    if rating_col and is_structured:
        df["label"] = df[rating_col].apply(normalize_label)
        text_only_mode = False
    else:
        df["label"] = df[review_col].apply(heuristic_sentiment)
        rating_col = "label"
        text_only_mode = True

    keep = [c for c in [product_col, review_col, rating_col, "label"] if c and c in df.columns]
    df = df[list(dict.fromkeys(keep))].copy()
    df = df.dropna(subset=[review_col])
    df["cleaned_review"] = df[review_col].apply(clean_text)
    df = df[df["cleaned_review"].str.strip().astype(bool)]

    label_counts = df["label"].value_counts()
    valid_labels = label_counts[label_counts >= 2].index.tolist()
    df = df[df["label"].isin(valid_labels)]

    if len(df) < 6:
        return {"error": (
            f"Not enough data to train the model ({len(df)} usable rows found). "
            "For DOCX/PDF/TXT files the system analyses every paragraph/sentence. "
            "Try a longer document or use a CSV with 'review' and 'rating' columns."
        )}

    X = df["cleaned_review"]
    y = df["label"]

    test_size = 0.2 if len(df) >= 20 else 0.15
    try:
        X_train, X_test, y_train, y_test = train_test_split(
            X, y, test_size=test_size, random_state=42,
            stratify=y if len(valid_labels) > 1 else None
        )
    except ValueError:
        X_train, X_test, y_train, y_test = train_test_split(
            X, y, test_size=test_size, random_state=42
        )

    pipeline = Pipeline([
        ("tfidf", TfidfVectorizer(max_features=8000, ngram_range=(1, 2))),
        ("clf", LogisticRegression(max_iter=1000, class_weight="balanced")),
    ])
    pipeline.fit(X_train, y_train)

    y_pred = pipeline.predict(X_test)
    accuracy = round(accuracy_score(y_test, y_pred) * 100, 1)
    report = classification_report(y_test, y_pred, output_dict=True, zero_division=0)

    df["predicted"] = pipeline.predict(df["cleaned_review"])
    proba = pipeline.predict_proba(df["cleaned_review"])
    classes = list(pipeline.classes_)
    df["confidence"] = (np.max(proba, axis=1) * 100).round(1)

    for cls in ["Positive", "Negative", "Neutral"]:
        key = cls.lower() + "_score"
        df[key] = (proba[:, classes.index(cls)] * 100).round(1) if cls in classes else 0.0

    counts = df["predicted"].value_counts().to_dict()
    per_class = {}
    for cls in ["Positive", "Negative", "Neutral"]:
        if cls in report:
            per_class[cls] = {
                "precision": round(report[cls]["precision"] * 100, 1),
                "recall": round(report[cls]["recall"] * 100, 1),
                "f1": round(report[cls]["f1-score"] * 100, 1),
                "support": int(report[cls]["support"]),
            }

    orig_rating_col = "label" if text_only_mode else rating_col
    display = df.rename(columns={
        review_col: "review",
        "positive_score": "pos_score",
        "negative_score": "neg_score",
        "neutral_score": "neu_score",
    }).copy()
    display["rating"] = "—" if text_only_mode else display[orig_rating_col].astype(str)
    display["actual"] = display["label"].astype(str)
    display["review"] = display["review"].astype(str).str.slice(0, 250)
    keep_cols = ["review", "rating", "actual", "predicted", "confidence",
                 "pos_score", "neg_score", "neu_score"]
    if product_col and product_col in display.columns:
        display["product"] = display[product_col].astype(str).str.slice(0, 80)
        keep_cols = ["product"] + keep_cols
    display = display[keep_cols].reset_index(drop=True)

    return {
        "accuracy": accuracy,
        "total": len(df),
        "counts": counts,
        "per_class": per_class,
        "table": display,
        "has_product": product_col is not None,
        "text_only": text_only_mode,
        "file_type": ext,
    }

# ────────────────────────────────────────────────────────────────
# Theme
# ────────────────────────────────────────────────────────────────
import html as _html

LIGHT = dict(
    bg="#f8fafc", card="#ffffff", text="#0f172a", subtext="#64748b",
    border="#e2e8f0", shadow="rgba(15,23,42,0.06)",
    accent="#6366f1", accent_soft="#eef2ff",
    green="#16a34a", green_soft="#dcfce7",
    red="#dc2626", red_soft="#fee2e2",
    amber="#d97706", amber_soft="#fef3c7",
    gray="#64748b", gray_soft="#f1f5f9",
    blue="#2563eb", blue_soft="#dbeafe",
)
DARK = dict(
    bg="#0b1220", card="#131c2e", text="#e5e7eb", subtext="#94a3b8",
    border="#243046", shadow="rgba(0,0,0,0.35)",
    accent="#818cf8", accent_soft="rgba(129,140,248,0.15)",
    green="#4ade80", green_soft="rgba(34,197,94,0.16)",
    red="#f87171", red_soft="rgba(239,68,68,0.16)",
    amber="#fbbf24", amber_soft="rgba(217,119,6,0.18)",
    gray="#9ca3af", gray_soft="rgba(148,163,184,0.14)",
    blue="#60a5fa", blue_soft="rgba(37,99,235,0.16)",
)

if "theme" not in st.session_state:
    st.session_state.theme = "light"
if "result" not in st.session_state:
    st.session_state.result = None

T = DARK if st.session_state.theme == "dark" else LIGHT


def esc(x):
    return _html.escape(str(x))


# ────────────────────────────────────────────────────────────────
# Global CSS
# ────────────────────────────────────────────────────────────────
st.markdown(f"""
<style>
.stApp {{
    background: {T['bg']};
    color: {T['text']};
}}
[data-testid="stHeader"] {{ background: transparent; }}
h1, h2, h3, h4, p, span, label, .stMarkdown {{ color: {T['text']}; }}
[data-testid="stCaptionContainer"], .stCaption {{ color: {T['subtext']} !important; }}

/* buttons */
.stButton>button {{
    background: linear-gradient(135deg, {T['accent']}, #8b5cf6);
    color: white; border: none; border-radius: 10px; font-weight: 600;
    padding: 0.55rem 1.2rem; box-shadow: 0 2px 6px {T['shadow']};
}}
.stButton>button:disabled {{
    background: {T['gray_soft']}; color: {T['subtext']};
}}
.stDownloadButton>button {{
    background: {T['card']}; color: {T['text']}; border: 1px solid {T['border']};
    border-radius: 10px; font-weight: 600;
}}

/* text area / inputs */
.stTextArea textarea, .stSelectbox div[data-baseweb="select"] > div {{
    background: {T['card']} !important; color: {T['text']} !important;
    border-radius: 10px !important; border: 1px solid {T['border']} !important;
}}

/* file uploader */
[data-testid="stFileUploaderDropzone"] {{
    background: {T['accent_soft']} !important;
    border: 2px dashed {T['accent']} !important;
    border-radius: 14px !important;
}}
[data-testid="stFileUploaderDropzoneInstructions"] svg {{ color: {T['accent']}; }}

hr {{ border-color: {T['border']}; }}

/* card */
.sml-card {{
    background: {T['card']}; border: 1px solid {T['border']};
    border-radius: 16px; padding: 20px 22px; margin-bottom: 6px;
    box-shadow: 0 1px 3px {T['shadow']};
}}
.sml-card-title {{ font-weight: 700; font-size: 1.02rem; margin-bottom: 2px; }}
.sml-card-desc {{ color: {T['subtext']}; font-size: 0.85rem; margin-bottom: 10px; }}

/* nav */
.sml-nav {{
    display:flex; align-items:center; justify-content:space-between;
    padding: 6px 2px 18px 2px;
}}
.sml-logo {{ display:flex; align-items:center; gap:10px; }}
.sml-logo-icon {{ font-size:1.8rem; }}
.sml-logo-title {{ font-weight:800; font-size:1.25rem; line-height:1.1; }}
.sml-logo-sub {{ color:{T['subtext']}; font-size:0.78rem; }}
.sml-badge {{
    display:inline-block; background:{T['accent_soft']}; color:{T['accent']};
    font-size:0.7rem; font-weight:700; letter-spacing:0.04em;
    padding:4px 10px; border-radius:999px;
}}

/* hero */
.sml-hero-title {{ font-size:1.9rem; font-weight:800; margin-bottom:4px; }}
.sml-hero-sub {{ color:{T['subtext']}; font-size:0.95rem; max-width:780px; }}

/* format pills */
.fmt-row {{ display:flex; gap:8px; flex-wrap:wrap; margin-top:14px; }}
.fmt-pill {{
    display:flex; align-items:center; gap:6px; font-size:0.8rem; font-weight:600;
    padding:6px 12px; border-radius:999px; border:1px solid transparent;
}}

/* step label */
.step-label {{ display:flex; align-items:center; gap:8px; font-weight:700; margin:22px 0 10px 0; }}
.step-num {{
    display:inline-flex; align-items:center; justify-content:center;
    width:22px; height:22px; border-radius:50%; background:{T['accent']};
    color:white; font-size:0.75rem; font-weight:800;
}}

/* alert */
.alert-warn {{
    background:{T['amber_soft']}; color:{T['amber']}; border:1px solid {T['amber']}33;
    border-radius:10px; padding:12px 14px; font-size:0.88rem; margin-bottom:14px;
}}

/* metric cards */
.metric-row {{ display:flex; gap:14px; flex-wrap:wrap; margin-bottom: 20px; }}
.metric-card {{
    position:relative; flex:1; min-width:150px; background:{T['card']};
    border:1px solid {T['border']}; border-radius:14px; padding:16px 18px;
    box-shadow:0 1px 3px {T['shadow']}; border-top:3px solid var(--mc);
}}
.metric-label {{ font-size:0.72rem; font-weight:700; color:{T['subtext']}; letter-spacing:0.03em; text-transform:uppercase; }}
.metric-value {{ font-size:1.7rem; font-weight:800; margin-top:6px; color:var(--mc); }}
.metric-icon {{ position:absolute; top:14px; right:16px; font-size:1.35rem; opacity:0.4; }}

/* badges (table) */
.badge {{ display:inline-block; padding:2px 10px; border-radius:999px; font-size:0.76rem; font-weight:700; }}
.badge-pos {{ background:{T['green_soft']}; color:{T['green']}; }}
.badge-neg {{ background:{T['red_soft']}; color:{T['red']}; }}
.badge-neu {{ background:{T['amber_soft']}; color:{T['amber']}; }}
.badge-gray {{ background:{T['gray_soft']}; color:{T['gray']}; }}

/* table */
.table-wrap {{ max-height:520px; overflow:auto; border:1px solid {T['border']}; border-radius:12px; }}
table.sml-table {{ width:100%; border-collapse:collapse; font-size:0.84rem; }}
table.sml-table thead th {{
    position:sticky; top:0; background:{T['card']}; color:{T['subtext']};
    text-align:left; padding:10px 12px; font-size:0.72rem; text-transform:uppercase;
    letter-spacing:0.03em; border-bottom:1px solid {T['border']}; z-index:1;
}}
table.sml-table tbody td {{
    padding:9px 12px; border-bottom:1px solid {T['border']}; vertical-align:middle;
    color:{T['text']};
}}
table.sml-table tbody tr:hover {{ background:{T['accent_soft']}; }}
.review-cell {{ max-width:340px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }}
.conf-wrap {{ display:flex; align-items:center; gap:8px; min-width:110px; }}
.conf-track {{ flex:1; height:6px; border-radius:999px; background:{T['gray_soft']}; overflow:hidden; }}
.conf-fill {{ height:100%; border-radius:999px; }}
.match-ok {{ color:{T['green']}; font-weight:800; }}
.match-bad {{ color:{T['red']}; font-weight:800; }}
</style>
""", unsafe_allow_html=True)


# ────────────────────────────────────────────────────────────────
# Chart helpers (theme-aware)
# ────────────────────────────────────────────────────────────────
CLASS_COLOR = {"Positive": T["green"], "Negative": T["red"], "Neutral": T["amber"]}


def sentiment_bar_chart(pct_map):
    labels = list(pct_map.keys())
    fig = go.Figure(go.Bar(
        x=list(pct_map.values()), y=labels, orientation="h",
        marker_color=[CLASS_COLOR.get(l, T["gray"]) for l in labels],
        text=[f"{v}%" for v in pct_map.values()], textposition="outside",
    ))
    fig.update_layout(
        xaxis_range=[0, 100], height=180, margin=dict(l=10, r=30, t=10, b=10),
        showlegend=False, paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color=T["text"]),
        xaxis=dict(showgrid=False, color=T["text"], tickfont=dict(color=T["text"], size=12),
                   linecolor=T["border"]),
        yaxis=dict(color=T["text"], tickfont=dict(color=T["text"], size=13),
                   linecolor=T["border"]),
    )
    return fig


def pie_chart(counts):
    labels = list(counts.keys())
    values = list(counts.values())
    colors = [CLASS_COLOR.get(l, T["gray"]) for l in labels]
    fig = go.Figure(go.Pie(labels=labels, values=values, marker_colors=colors, hole=0.55,
                            textinfo="percent", textfont=dict(color=T["text"])))
    fig.update_layout(
        margin=dict(l=10, r=10, t=10, b=10), height=300,
        paper_bgcolor="rgba(0,0,0,0)",
        legend=dict(orientation="h", y=-0.1, font=dict(color=T["text"])),
    )
    return fig


def f1_chart(per_class):
    labels = list(per_class.keys())
    f1s = [per_class[l]["f1"] for l in labels]
    colors = [CLASS_COLOR.get(l, T["gray"]) for l in labels]
    fig = go.Figure(go.Bar(
        x=labels, y=f1s, marker_color=colors, width=0.5,
        text=[f"{v}%" for v in f1s], textposition="outside",
    ))
    fig.update_layout(
        yaxis_range=[0, 115], height=300, margin=dict(l=10, r=10, t=10, b=10),
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color=T["text"]),
        yaxis=dict(showgrid=True, gridcolor=T["border"], color=T["text"],
                   tickfont=dict(color=T["text"], size=12), linecolor=T["border"]),
        xaxis=dict(color=T["text"], tickfont=dict(color=T["text"], size=13),
                   linecolor=T["border"]),
    )
    return fig


def badge(label):
    cls = {"Positive": "badge-pos", "Negative": "badge-neg", "Neutral": "badge-neu"}.get(label, "badge-gray")
    return f'<span class="badge {cls}">{esc(label)}</span>'


def confidence_bar(val):
    color = T["accent"]
    return (f'<div class="conf-wrap"><div class="conf-track">'
            f'<div class="conf-fill" style="width:{val}%;background:{color};"></div>'
            f'</div><span>{val:.1f}%</span></div>')


def build_table_html(df_page, has_product, start_idx):
    cols_html = "<th>#</th>"
    if has_product:
        cols_html += "<th>Product</th>"
    cols_html += "<th>Review</th><th>Rating</th><th>Actual</th><th>Predicted</th><th>Match</th><th>Confidence</th>"

    rows_html = []
    for i, (_, row) in enumerate(df_page.iterrows()):
        match_ok = row["actual"] == row["predicted"]
        match_html = '<span class="match-ok">✓</span>' if match_ok else '<span class="match-bad">✗</span>'
        product_html = f"<td>{esc(row['product'])}</td>" if has_product else ""
        rows_html.append(
            "<tr>"
            f"<td>{start_idx + i + 1}</td>"
            f"{product_html}"
            f"<td class='review-cell' title='{esc(row['review'])}'>{esc(row['review'])}</td>"
            f"<td>{esc(row['rating'])}</td>"
            f"<td>{badge(row['actual'])}</td>"
            f"<td>{badge(row['predicted'])}</td>"
            f"<td>{match_html}</td>"
            f"<td>{confidence_bar(row['confidence'])}</td>"
            "</tr>"
        )

    return (
        '<div class="table-wrap"><table class="sml-table">'
        f"<thead><tr>{cols_html}</tr></thead>"
        f"<tbody>{''.join(rows_html)}</tbody>"
        "</table></div>"
    )


# ────────────────────────────────────────────────────────────────
# Nav + theme toggle
# ────────────────────────────────────────────────────────────────
nav_l, nav_r = st.columns([6, 1.4])
with nav_l:
    st.markdown(
        f"""
        <div class="sml-logo">
          <span class="sml-logo-icon">🧠</span>
          <div>
            <div class="sml-logo-title">SentimentML</div>
            <div class="sml-logo-sub">TF-IDF · Logistic Regression · NLP</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
with nav_r:
    st.markdown('<div style="text-align:right;"><span class="sml-badge">ML POWERED</span></div>',
                unsafe_allow_html=True)
    dark_on = st.toggle("🌙 Dark mode", value=(st.session_state.theme == "dark"), key="theme_switch")
    new_theme = "dark" if dark_on else "light"
    if new_theme != st.session_state.theme:
        st.session_state.theme = new_theme
        st.rerun()

st.markdown('<div class="sml-hero-title">Sentiment Analysis Dashboard</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="sml-hero-sub">Upload any document — the model extracts text, trains on your data '
    'using NLP + ML, and predicts sentiment as <strong>Positive</strong>, <strong>Negative</strong>, '
    'or <strong>Neutral</strong>.</div>',
    unsafe_allow_html=True,
)

FORMAT_PILLS = [
    ("📊", "CSV", T["green"], T["green_soft"]),
    ("📄", "PDF", T["red"], T["red_soft"]),
    ("📝", "DOCX", T["blue"], T["blue_soft"]),
    ("📈", "XLSX", T["green"], T["green_soft"]),
    ("📃", "TXT", T["gray"], T["gray_soft"]),
    ("🔣", "JSON", T["amber"], T["amber_soft"]),
]
pills_html = "".join(
    f'<span class="fmt-pill" style="color:{c};background:{bg};">{icon} {label}</span>'
    for icon, label, c, bg in FORMAT_PILLS
)
st.markdown(f'<div class="fmt-row">{pills_html}</div>', unsafe_allow_html=True)

st.markdown("<div style='height:22px'></div>", unsafe_allow_html=True)

# ────────────────────────────────────────────────────────────────
# Live text analysis
# ────────────────────────────────────────────────────────────────
st.markdown('<div class="step-label"><span class="step-num">✍</span> Analyse any text instantly</div>',
            unsafe_allow_html=True)

with st.container(border=False):
    st.markdown('<div class="sml-card">', unsafe_allow_html=True)
    st.markdown('<div class="sml-card-title">Type or paste your text</div>', unsafe_allow_html=True)
    st.markdown('<div class="sml-card-desc">Reviews, sentences, paragraphs — anything. '
                'Get instant sentiment predictions.</div>', unsafe_allow_html=True)

    live_text = st.text_area(
        "Type or paste your text", label_visibility="collapsed",
        placeholder="e.g. The product quality was amazing and delivery was super fast! Highly recommend.",
        height=110, key="live_text_input",
    )
    word_count = len(re.findall(r"\b\w+\b", live_text)) if live_text else 0
    st.caption(f"{word_count} words")

    analyse_clicked = st.button("🔍 Analyse Sentiment", disabled=word_count == 0, key="analyse_btn")

    if analyse_clicked:
        if len(live_text.strip()) < 3:
            st.markdown('<div class="alert-warn">Text is too short to analyse.</div>', unsafe_allow_html=True)
        else:
            st.session_state["live_result"] = predict_live_text(live_text)

    live_result = st.session_state.get("live_result")
    if live_result:
        c1, c2 = st.columns([1, 2])
        with c1:
            st.markdown(
                f'<div style="font-size:2.4rem;">{live_result["emoji"]}</div>'
                f'<div style="font-weight:800;font-size:1.15rem;">{live_result["overall"]}</div>'
                f'<div style="color:{T["subtext"]};font-size:0.8rem;">Overall sentiment</div>',
                unsafe_allow_html=True,
            )
        with c2:
            st.plotly_chart(
                sentiment_bar_chart({
                    "Positive": live_result["pos_pct"],
                    "Negative": live_result["neg_pct"],
                    "Neutral": live_result["neu_pct"],
                }),
                width='stretch', key="live_bar_chart",
            )
        if len(live_result["sentences"]) > 1:
            with st.expander("Sentence-level breakdown"):
                for s in live_result["sentences"]:
                    st.markdown(f"{badge(s['label'])} &nbsp; {esc(s['sentence'])}", unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

st.markdown(
    f'<div style="text-align:center;color:{T["subtext"]};margin:18px 0;font-size:0.85rem;">'
    "— or upload a full dataset —</div>",
    unsafe_allow_html=True,
)

# ────────────────────────────────────────────────────────────────
# File upload + train
# ────────────────────────────────────────────────────────────────
st.markdown('<div class="step-label"><span class="step-num">1</span> Upload your file</div>',
            unsafe_allow_html=True)

st.markdown('<div class="sml-card">', unsafe_allow_html=True)
uploaded = st.file_uploader(
    "Drop your file here or click to browse",
    type=sorted(ALLOWED_EXTENSIONS), key="file_uploader",
)

st.markdown('<div class="step-label"><span class="step-num">2</span> Train &amp; Predict</div>',
            unsafe_allow_html=True)
run_clicked = st.button("⚡ Train Model & Analyse Sentiment", disabled=uploaded is None, key="train_btn")
st.markdown("</div>", unsafe_allow_html=True)

if run_clicked and uploaded is not None:
    with st.spinner("Extracting text & training model… usually takes 5–20 seconds"):
        file_bytes = uploaded.read()
        st.session_state["result"] = run_pipeline(uploaded.name, file_bytes)
        st.session_state["table_page"] = 0

# ────────────────────────────────────────────────────────────────
# Results
# ────────────────────────────────────────────────────────────────
result = st.session_state.get("result")
if result:
    if "error" in result:
        st.markdown(f'<div class="alert-warn">{esc(result["error"])}</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div class="step-label"><span class="step-num">3</span> Results</div>',
                    unsafe_allow_html=True)

        if result["text_only"]:
            st.markdown(
                '<div class="alert-warn">⚠️ No rating column found — sentiment was predicted using '
                'keyword-based heuristics. For best accuracy, use a CSV/XLSX/JSON with a '
                '<strong>rating</strong> or <strong>sentiment</strong> column.</div>',
                unsafe_allow_html=True,
            )

        counts = result["counts"]
        cards = [
            ("MODEL ACCURACY", f"{result['accuracy']}%", "🎯", T["accent"]),
            ("POSITIVE", str(counts.get("Positive", 0)), "😊", T["green"]),
            ("NEGATIVE", str(counts.get("Negative", 0)), "😞", T["red"]),
            ("NEUTRAL", str(counts.get("Neutral", 0)), "😐", T["amber"]),
            ("TOTAL REVIEWED", str(result["total"]), "📊", T["gray"]),
        ]
        cards_html = "".join(
            f'<div class="metric-card" style="--mc:{color};">'
            f'<div class="metric-label">{label}</div>'
            f'<div class="metric-value">{value}</div>'
            f'<div class="metric-icon">{icon}</div>'
            "</div>"
            for label, value, icon, color in cards
        )
        st.markdown(f'<div class="metric-row">{cards_html}</div>', unsafe_allow_html=True)

        chart_col1, chart_col2 = st.columns(2)
        with chart_col1:
            st.markdown('<div class="sml-card">', unsafe_allow_html=True)
            st.markdown('<div class="sml-card-title">Sentiment Distribution</div>', unsafe_allow_html=True)
            st.markdown('<div class="sml-card-desc">Breakdown of all predictions by class</div>',
                        unsafe_allow_html=True)
            st.plotly_chart(pie_chart(counts), width='stretch', key="pie_chart")
            st.markdown("</div>", unsafe_allow_html=True)
        with chart_col2:
            st.markdown('<div class="sml-card">', unsafe_allow_html=True)
            st.markdown('<div class="sml-card-title">Per-Class F1 Score</div>', unsafe_allow_html=True)
            st.markdown('<div class="sml-card-desc">Balance of precision &amp; recall per sentiment class</div>',
                        unsafe_allow_html=True)
            if result["per_class"]:
                st.plotly_chart(f1_chart(result["per_class"]), width='stretch', key="f1_chart")
            else:
                st.info("Not enough class variety to compute per-class F1 scores.")
            st.markdown("</div>", unsafe_allow_html=True)

        # ── Predictions table ──────────────────────────────────
        st.markdown('<div class="sml-card">', unsafe_allow_html=True)
        st.markdown('<div class="sml-card-title">All Predictions</div>', unsafe_allow_html=True)

        table = result["table"]

        if "table_filter" not in st.session_state:
            st.session_state.table_filter = "All"

        FILTER_OPTIONS = ["All", "Positive", "Negative", "Neutral"]
        FILTER_COLOR = {"All": T["accent"], "Positive": T["green"],
                        "Negative": T["red"], "Neutral": T["amber"]}

        filter_css_rules = []
        for opt in FILTER_OPTIONS:
            active = st.session_state.table_filter == opt
            color = FILTER_COLOR[opt]
            if active:
                bg, fg, border = color, "#ffffff", color
            else:
                bg, fg, border = T["card"], T["text"], T["border"]
            filter_css_rules.append(f"""
            .st-key-filter_btn_{opt} button {{
                background: {bg} !important;
                color: {fg} !important;
                border: 1.5px solid {border} !important;
                border-radius: 999px !important;
                font-weight: 700 !important;
                padding: 0.3rem 1rem !important;
                box-shadow: none !important;
            }}""")
        st.markdown(f"<style>{''.join(filter_css_rules)}</style>", unsafe_allow_html=True)

        fcols = st.columns([0.8, 1.2, 1.3, 1.2, 5])
        for col, opt in zip(fcols, FILTER_OPTIONS):
            with col:
                if st.button(opt, key=f"filter_btn_{opt}"):
                    st.session_state.table_filter = opt
                    st.session_state.table_page = 0
                    st.rerun()

        filter_choice = st.session_state.table_filter
        shown = table if filter_choice == "All" else table[table["predicted"] == filter_choice]
        shown = shown.reset_index(drop=True)

        PAGE_SIZE = 50
        n_pages = max(1, (len(shown) - 1) // PAGE_SIZE + 1)
        page = st.session_state.get("table_page", 0)
        page = min(page, n_pages - 1)
        start = page * PAGE_SIZE
        page_df = shown.iloc[start:start + PAGE_SIZE]

        st.caption(f"Showing {min(start + 1, len(shown))}–{min(start + PAGE_SIZE, len(shown))} "
                   f"of {len(shown)} rows")

        st.markdown(
            build_table_html(page_df, result["has_product"], start),
            unsafe_allow_html=True,
        )

        if n_pages > 1:
            pcol1, pcol2, pcol3 = st.columns([1, 2, 1])
            with pcol1:
                if st.button("◀ Prev", disabled=page == 0, key="prev_page"):
                    st.session_state["table_page"] = page - 1
                    st.rerun()
            with pcol2:
                st.markdown(
                    f'<div style="text-align:center;color:{T["subtext"]};padding-top:8px;">'
                    f"Page {page + 1} of {n_pages}</div>",
                    unsafe_allow_html=True,
                )
            with pcol3:
                if st.button("Next ▶", disabled=page >= n_pages - 1, key="next_page"):
                    st.session_state["table_page"] = page + 1
                    st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)

        csv_bytes = table.to_csv(index=False).encode("utf-8")
        st.download_button(
            "⬇️ Download full results as CSV",
            data=csv_bytes, file_name="sentiment_results.csv", mime="text/csv",
        )
